"""
Диагностический скрипт.  Запуск из папки quiz_bot:

    python debug_parser.py путь/к/файлу.pdf
    python debug_parser.py путь/к/файлу.docx
    python debug_parser.py путь/к/файлу.txt
"""
import sys
import os
sys.stdout.reconfigure(encoding='utf-8')

if len(sys.argv) < 2:
    print("Использование: python debug_parser.py <путь_к_файлу>")
    sys.exit(1)

# Запускаем из папки quiz_bot, чтобы импорты работали
os.chdir(os.path.dirname(os.path.abspath(__file__)))

path = sys.argv[1]
ext  = os.path.splitext(path)[1].lower()
SEP  = '=' * 62

print(f"\n{SEP}")
print(f"Файл : {path}")
print(f"Тип  : {ext}")
print(SEP)


# ── PDF ───────────────────────────────────────────────────────────────────────
if ext == '.pdf':
    import pdfplumber

    with pdfplumber.open(path) as pdf:
        n_pages = len(pdf.pages)
        print(f"\nСтраниц: {n_pages}")

        # Count annotations
        total_annots = hl_annots = 0
        for page in pdf.pages:
            for a in (page.annots or []):
                total_annots += 1
                if 'Highlight' in str(a.get('data', {}).get('Subtype', '')):
                    hl_annots += 1

        # Count yellow fill rects
        yellow_rects = 0
        for page in pdf.pages:
            for r in (page.rects or []):
                c = r.get('non_stroking_color')
                if isinstance(c, (list, tuple)):
                    if len(c) == 3 and c[0] > 0.65 and c[1] > 0.65 and c[2] < 0.35:
                        yellow_rects += 1
                    elif len(c) == 4 and c[2] > 0.65 and c[3] < 0.25:
                        yellow_rects += 1

        # Count yellow chars
        yellow_chars = 0
        for page in pdf.pages:
            for ch in (page.chars or []):
                c = ch.get('non_stroking_color')
                if isinstance(c, (list, tuple)):
                    if len(c) == 3 and c[0] > 0.65 and c[1] > 0.65 and c[2] < 0.35:
                        yellow_chars += 1
                    elif len(c) == 4 and c[2] > 0.65 and c[3] < 0.25:
                        yellow_chars += 1

        print(f"Аннотаций всего          : {total_annots}")
        print(f"  из них /Highlight      : {hl_annots}")
        print(f"Жёлтых прямоугольников  : {yellow_rects}")
        print(f"Символов с жёлтым фоном : {yellow_chars}")

        print("\n─── Текст стр. 1 (первые 1 000 символов) ───")
        print((pdf.pages[0].extract_text() or '')[:1000])

    print()
    from parsers.pdf_parser import _collect_highlights, _parse_with_highlights, _parse_plain

    highlights = _collect_highlights(path)
    print(f"Highlight-фрагментов найдено: {len(highlights)}")
    for h in highlights[:8]:
        print(f"  → '{h[:90]}'")

    print()
    qs_hl = _parse_with_highlights(path)
    print(f"Вопросов через highlight-детект : {len(qs_hl)}")
    for q in qs_hl[:3]:
        print(f"  Q: {q['text'][:65]}…")
        print(f"     Правильный [{q['correct']}]: {q['options'][q['correct']][:50]}")

    qs_plain = _parse_plain(path)
    print(f"\nВопросов через plain-text       : {len(qs_plain)}")


# ── DOCX ──────────────────────────────────────────────────────────────────────
elif ext == '.docx':
    from docx import Document
    doc = Document(path)
    print(f"\nАбзацев: {len(doc.paragraphs)}")

    try:
        from docx.enum.text import WD_COLOR_INDEX
        _COLORS = {WD_COLOR_INDEX.YELLOW, WD_COLOR_INDEX.BRIGHT_GREEN, WD_COLOR_INDEX.CYAN}
        hl_runs = [
            (p.text[:40], run.text)
            for p in doc.paragraphs
            for run in p.runs
            if run.font.highlight_color in _COLORS and run.text.strip()
        ]
        print(f"Выделенных (цветных) runs: {len(hl_runs)}")
        for par, run in hl_runs[:5]:
            print(f"  В абз «{par}…» → '{run}'")
    except Exception as e:
        print(f"Ошибка при проверке highlight: {e}")

    from parsers.docx_parser import parse_docx
    qs = parse_docx(path)
    print(f"\nВопросов найдено: {len(qs)}")
    for q in qs[:3]:
        print(f"  Q: {q['text'][:60]}  → correct={q['correct']}")


# ── TXT ───────────────────────────────────────────────────────────────────────
elif ext == '.txt':
    with open(path, encoding='utf-8') as f:
        text = f.read()
    print(f"\nСимволов: {len(text)}")
    print("\nПервые 600 символов:")
    print(text[:600])

    from parsers.txt_parser import parse_txt
    qs = parse_txt(text)
    print(f"\nВопросов найдено: {len(qs)}")
    for q in qs[:3]:
        print(f"  Q: {q['text'][:60]}  → correct={q['correct']}")

else:
    print(f"Формат {ext!r} не поддерживается этим скриптом.")
    sys.exit(1)

print(f"\n{SEP}")
print("Готово. Скопируй вывод, если нужна помощь с диагностикой.")
print(SEP)
