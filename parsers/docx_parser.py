import re
from docx import Document
from parsers.base import texts_similar, strip_option_prefix
from parsers.txt_parser import parse_txt, _parse_block_no_answer


def parse_docx(file_path: str) -> list[dict]:
    questions = _parse_with_highlights(file_path)
    if questions:
        return questions
    return parse_txt(_extract_text(file_path))


# ── Highlight pipeline ────────────────────────────────────────────────────────

def _parse_with_highlights(file_path: str) -> list[dict]:
    try:
        from docx.enum.text import WD_COLOR_INDEX
        _COLORS = {WD_COLOR_INDEX.YELLOW, WD_COLOR_INDEX.BRIGHT_GREEN, WD_COLOR_INDEX.CYAN}

        doc = Document(file_path)
        highlights: set[str] = set()
        para_lines: list[str] = []

        for para in doc.paragraphs:
            para_lines.append(para.text.strip())
            for run in para.runs:
                if run.font.highlight_color in _COLORS and run.text.strip():
                    highlights.add(run.text.strip())

        if not highlights:
            return []

        full_text = '\n'.join(para_lines)
        raw: list[dict] = []
        for block in re.split(r'\n\s*\n', full_text):
            q = _parse_block_no_answer(block.strip())
            if q:
                raw.append(q)

        for q in raw:
            for i, opt in enumerate(q['options']):
                opt_body = strip_option_prefix(opt)
                if any(texts_similar(opt_body, hl) for hl in highlights):
                    q['correct'] = i
                    break

        return [q for q in raw if q.get('correct') is not None]

    except Exception:
        return []


# ── Plain-text extraction ─────────────────────────────────────────────────────

def _extract_text(file_path: str) -> str:
    doc = Document(file_path)
    return '\n'.join(p.text.strip() for p in doc.paragraphs if p.text.strip())
